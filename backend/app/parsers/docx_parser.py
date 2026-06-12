"""DOCX text extractor using python-docx."""

from io import BytesIO

from docx import Document as DocxDocument

from app.parsers.base import BaseParser
from app.utils.logger import get_logger

logger = get_logger(__name__)


class DocxParser(BaseParser):
    """Extract text from .docx files using python-docx."""

    def parse(self, file_bytes: bytes, filename: str = "") -> str:
        """
        Open a DOCX document from bytes and join all paragraph texts.

        Args:
            file_bytes: Raw DOCX bytes.
            filename:   Original filename (informational).

        Returns:
            All paragraph text joined by newlines.
        """
        try:
            doc = DocxDocument(BytesIO(file_bytes))
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text.strip())
            extracted = "\n".join(paragraphs)
            logger.info(
                f"DOCX parsed: {filename!r} — {len(paragraphs)} paragraphs, "
                f"{len(extracted)} chars."
            )
            return extracted
        except Exception as exc:
            logger.error(f"DocxParser failed for {filename!r}: {exc}")
            raise
